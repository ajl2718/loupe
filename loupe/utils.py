# functions for detecting and extracting PII
import pandas as pd
import numpy as np
import spacy
from copy import deepcopy
from time import time
from spacy.matcher import Matcher
from math import ceil, log
from docx import Document
import pdfplumber

# the default types of PII to look for (i.e., all of them)
DEFAULT_PII_TYPES = [
    "PERSON_NAME",
    "PHONE_AU",
    "EMAIL",
    "URL",
    "MEDICARE_NUMBER",
    "TAX_FILE_NUMBER",
    "POSTCODE",
    "DATE",
    "CAR_REG_VIC",
    "CREDIT_CARD_NUMBER",
    "NORP",
    "GPE",
    "LANGUAGE",
    "ORG",
    "FAC",
    "MONEY",
]

ENTITY_TYPES = [
    "PERSON",
    "PHONE_AU",
    "EMAIL",
    "URL",
    "MEDICARE_CANDIDATE",
    "TFN_CANDIDATE",
    "POSTCODE",
    "DATE",
    "CAR_REG_VIC",
    "CREDIT_CARD_CANDIDATE",
    "NORP",
    "GPE",
    "LANGUAGE",
    "ORG",
    "FAC",
    "MONEY",
]

ENTITY_PII_MAPPING = dict(zip(DEFAULT_PII_TYPES, ENTITY_TYPES))

no_check_digit = lambda x: True


def get_text_from_pdf(pdf):
    """
    Given a pdfplumber PDF object, extract all the text and return a list of strs

    Args:
    pdf: pdfplumber pdf object

    Return:
    output: list of strs with a str corresponding to each page of the PDF
    """
    output = []
    for page in pdf.pages:
        output.append(page.extract_text())

    return output


def get_table_text(table):
    """
    Given a docx Table object, extract all the text from it

    Args:
    table: docx Table object

    Return:
    table_texts: list of strs containing the text in table
    """
    table_texts = []

    for row in table.rows:
        for cell in row.cells:
            # When a newline appears, split into separate texts as we assume these are separate
            # bits of text
            table_texts += cell.text.split("\n")

    return table_texts


## TODO: investigate other python libraries for dealing with word documents
## Does this extract all the text? Need to check this?
## Does this properly extract text, e.g., not cutting off parts of text.
def get_text_from_worddoc(doc):
    """
    Extract all the text (from paragraphs, tables, footnotes, headers, titles...)
    from a Word document. Return a list of strings.

    Args:
    doc: a docx Document object

    Return:
    doc_texts: list of strs with the text from worddoc

    """
    # get all the text from paragraphs
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]

    # get all the text from table cells
    cells = [text for table in doc.tables for text in get_table_text(table)]

    # repeat these for the headers and footers
    headers = [section.header for section in doc.sections]
    footers = [section.footer for section in doc.sections]

    header_paragraphs = [
        paragraph.text for header in headers for paragraph in header.paragraphs
    ]
    footer_paragraphs = [
        paragraph.text for footer in footers for paragraph in footer.paragraphs
    ]

    header_cells = [
        text
        for header in headers
        for table in header.tables
        for text in get_table_text(table)
    ]
    footer_cells = [
        text
        for footer in footers
        for table in footer.tables
        for text in get_table_text(table)
    ]

    return (
        paragraphs
        + header_paragraphs
        + footer_paragraphs
        + cells
        + header_cells
        + footer_cells
    )


def process_text(text):
    """
    Process str to remove certain characters so that spacy can work with it
    """
    # remove the apostrophes, new lines, line breaks
    if text == None:
        return ""
    else:
        text = (
            text.replace("'", "")
            .replace("\n", " ")
            .replace("\r", "")
            .replace("\b", " ")
        )
        return text


def process_df(df):
    """
    Given a data frame, fill nan values with 0 or '' depending on datatype
    Ensures that column from dataframe can be parsed by pii detection functions
    """
    dtypes = df.dtypes

    for column in dtypes.index:
        if dtypes[column] == "O":
            df[column] = df[column].astype(str).fillna("")
        else:
            df[column] = df[column].fillna(0)

    return df


def string_is_integer(s):
    """
    Check if a string s represents an integer
    Returns the integer or False if an error is raised
    """
    try:
        return int(s.replace(" ", ""))
    except ValueError:
        return False


def get_digits(number: int):
    """
    Given an integer, return an array containing each of its digits

    Args:
    number: an integer

    Returns:
    list of ints: each entry is a digit of number
    """
    num_digits = ceil(log(number, 10))
    digits = []
    for n in range(0, num_digits):
        exponent = num_digits - 1 - n
        digit = number // (10 ** exponent) - 10 * (number // (10 ** (exponent + 1)))
        digits.append(digit)
    return digits


def is_creditcard(number: int):
    """
    Given a number, check that it is a valid credit card number
    Based on check digits algorithm in: https://www.gizmodo.com.au/2014/01/how-credit-card-numbers-work/
    Luhn algorithm

    Args:
    number: integer with the credit card number

    Returns:
    True if it is a valid credit card number and False otherwise
    """
    is_16_digits = (number < 10 ** 16) & (number > 10 ** 15)
    # mapping from doubling a digit and then summing the resulting digits if >= 10
    digit_mapping = {0: 0, 1: 2, 2: 4, 3: 6, 4: 8, 5: 1, 6: 3, 7: 5, 8: 7, 9: 9}
    if is_16_digits == False:
        return False
    else:
        digits = get_digits(number)
    # weighted sum over digits
    weighted_sum = sum([digit_mapping[digit] for digit in digits[0::2]]) + sum(
        digits[1::2]
    )
    return weighted_sum % 10 == 0


def is_tfn(number: int) -> bool:
    """
    Given a number, check that it is a valid Australian Tax File Number
    Based on check digits algorithm in: http://www.mathgen.ch/codes/tfn.html

    Args:
    number: integer with the candidate tax file number

    Returns:
    True if it is a valid TFN and False otherwise
    """
    weights = (1, 4, 3, 7, 5, 8, 6, 9, 10)

    is_9_digits = (number < 10 ** 9) & (number > 10 ** 8)

    if is_9_digits == False:
        return False
    else:
        digits = get_digits(number)

    weighted_sum = sum([weight * digit for weight, digit in zip(weights, digits)])

    return weighted_sum % 11 == 0


def is_medicare(number: int):
    """
    Given a number, check that it is a valid Medicare number
    Based on check digits algorithm in: https://www.clearwater.com.au/code/medicare

    Args:
    number: integer with the candidate medicare number

    Returns:
    True if it is a valid medicare number and False otherwise
    """
    weights = (1, 3, 7, 9, 1, 3, 7, 9)
    is_10_digits = (number < 10 ** 10) & (number > 10 ** 9)

    if is_10_digits == False:
        return False
    else:
        digits = get_digits(number)

    weighted_sum = sum([weight * digit for weight, digit in zip(weights, digits[:-2])])

    return weighted_sum % 10 == digits[-2]


def candidate_tfn_valid(tfn_candidate_span):
    """
    Given a spacy span representing a tax file number candidate check that it is
    a valid TFN.

    Args
    tfn_candidate_span: spacy span

    Return
    True if it is a TFN or False otherwise
    """
    tfn_candidate = string_is_integer(tfn_candidate_span.text.replace(" ", ""))
    if tfn_candidate:
        return is_tfn(tfn_candidate)
    else:
        return False


def candidate_medicare_valid(medicare_candidate_span):
    """
    Given a spacy span representing a medicare number candidate check that it is
    a valid medicare card number.

    Args
    medicare_candidate_span: spacy span

    Return
    True if it is a medicare card number or False otherwise
    """
    medicare_candidate = string_is_integer(
        medicare_candidate_span.text.replace(" ", "")
    )
    if medicare_candidate:
        return is_medicare(medicare_candidate)
    else:
        return False


def candidate_creditcard_valid(creditcard_candidate_span):
    """
    Given a spacy span representing a credit carrd number candidate check that it is
    a valid credit card number.

    Args
    creditcard_candidate_span: spacy span

    Return
    True if it is a credit card number or False otherwise
    """
    creditcard_candidate = string_is_integer(
        creditcard_candidate_span.text.replace(" ", "")
    )
    if creditcard_candidate:
        return is_creditcard(creditcard_candidate)
    else:
        return False


def is_nlp_type(pii_type):
    """
    checks a pii_type to see if it requires spaCy language model
    these include:
    - PERSON_NAME
    - GPE
    - ADDRESS
    - FAC
    - LOC
    - DATE
    """
    nlp_types = [
        "PERSON_NAME",
        "GPE",
        "ORG",
        "FAC",
        "LOC",
        "ADDRESS",
        "DATE",
        "NORP",
        "LANGUAGE",
    ]
    if pii_type in nlp_types:
        return True
    else:
        return False


# TODO: move the data, i.e., the patterns here to a YAML file (actually JSONL based on spacy documentation)
# Do we use the EntityRuler here?
def load_patterns(file):
    """
    Load up the rules for the patterns from jsonl file, as per spacy documentation.

    Args:
    file: path of jsonl file containing the patterns

    Returns:
    patterns: list of dicts of the form ("match_id": "NAME_OF_ENTITY", "patterns": [{patterns...}])
    """
    with open(file, "r") as f:
        outputs = f.readlines()

    outputs = [eval(output) for output in outputs]

    return outputs


# TODO: use on_match rules to tidy up detection of the entity spans
# TODO: use phrase matcher to extract rules defining various combinations of phone numbers (with spaces in between)
def create_matcher(nlp, pattern_file):
    """
    Create a spacy matcher object containing rules that define the entities we would
    like to detect. These are for well-structured entities such as postcodes, telephone numbers,
    medicare numbers.

    To do: move the rules to another function so that this is neater

    Args:
    nlp: a spacy language model
    pattern_file: str describing path to the .jsonl file with the patterns

    return: spacy.matcher object
    """
    matcher = Matcher(nlp.vocab)

    # load up the patterns from a file and then add each type of pattern
    patterns = load_patterns(pattern_file)
    for pattern in patterns:
        matcher.add(pattern["match_id"], pattern["patterns"])

    return matcher


# TODO: allow for different levels of masking: e.g., replace with entity type, hash, etc.
def replace_pii_in_text(doc, pii_spans):
    """
    Given a text and a list of entity spans corresponding to pii items replace each
    of elements of text corresponding to pii with '[PII_REMOVED]'

    Args
    doc: spacy Doc object of original text
    pii_spans: list of spacy.span objects describing the PII spans in text

    Return
    text_pii_removed: str with PII removed from original text in doc
    """

    # generate the text elements in doc
    original_text = [token.text for token in doc]

    # replace the pii elements with '[PII_REMOVED]'
    for pii_span in pii_spans:
        start, end = pii_span
        original_text[start:end] = ["****"]
    text_pii_removed = " ".join(original_text)

    return text_pii_removed


# to do: string processing to separate some characters that would make tokenization difficult
def redact_text(doc, pii_spans):
    """
    Given a text and a list of entity spans corresponding to pii items replace each
    of elements of text corresponding to pii with '[PII_REMOVED]'

    Args
    doc: spacy Doc object of original text
    pii_spans: list of spacy.span objects describing the PII spans in text

    Return
    text_pii_removed: str with PII removed from original text in doc
    """
    # generate the text elements in doc
    original_text = [token.text for token in doc]
    pii_spans.sort()
    if len(pii_spans) == 0:
        return " ".join(original_text)
    else:
        gap_start = 0
        gap_end = pii_spans[0][0]
        all_spans = [(gap_start, gap_end)]
        # create the text for the gaps and pii
        for n in range(0, len(pii_spans) - 1):
            gap_start = pii_spans[n][1]
            gap_end = pii_spans[n + 1][0]
            all_spans.append((pii_spans[n][0], pii_spans[n][1]))  # append pii_span
            all_spans.append((gap_start, gap_end))  # append gap_span
        all_spans.append((pii_spans[-1][0], pii_spans[-1][1]))
        all_spans.append((pii_spans[-1][1], -1))
        # odd number spans: keep the text
        # even number spans: redact
        redacted_list = []
        for n, span in enumerate(all_spans):
            if n % 2 == 0:
                redacted_list += original_text[span[0] : span[1]]
            else:
                redacted_list += ["***"]
        return " ".join(redacted_list)


def redact_text_new(text, pii_spans):
    """
    Given a text and a list of entity spans corresponding to pii items replace each
    of elements of text corresponding to pii with '[PII_REMOVED]'

    Args
    doc: spacy Doc object of original text
    pii_spans: list of spacy.span objects describing the PII spans in text

    Return
    text_pii_removed: str with PII removed from original text in doc
    """
    pii_spans.sort()
    if len(pii_spans) == 0:
        return text
    else:
        gap_start = 0
        gap_end = pii_spans[0][0]
        all_spans = [(gap_start, gap_end)]
        # create the text for the gaps and pii
        for n in range(0, len(pii_spans) - 1):
            gap_start = pii_spans[n][1]
            gap_end = pii_spans[n + 1][0]
            all_spans.append((pii_spans[n][0], pii_spans[n][1]))  # append pii_span
            all_spans.append((gap_start, gap_end))  # append gap_span
        all_spans.append((pii_spans[-1][0], pii_spans[-1][1]))
        all_spans.append((pii_spans[-1][1], -1))
        # odd number spans: keep the text
        # even number spans: redact
        redacted_list = []
        for n, span in enumerate(all_spans):
            if n % 2 == 0:
                redacted_list += text[span[0] : span[1]]
            else:
                redacted_list += ["***"]
        return "".join(redacted_list)


def redact_texts(texts, matcher, nlp, pii_types=DEFAULT_PII_TYPES):
    pii_with_spans = pii_from_texts_with_spans(texts, matcher, nlp, pii_types)
    redacted_list = []
    for n in range(0, len(texts)):
        text = texts[n]
        spans = []
        pii = pii_with_spans[n]
        for key in pii.keys():
            for pii_item in pii[key]:
                spans.append(pii_item[0])
        new_text = redact_text_new(text, spans)
        redacted_list.append(new_text)
    return redacted_list


# todo: apply additional tests for PERSON_NAME using dict of person names?
def pii_from_texts(
    texts,
    matcher,
    nlp,
    pii_types=DEFAULT_PII_TYPES,
    output="counts",
    random_sample="None",
    n_process=1,
    batch_size=2000,
):
    """
    Extract Personally Identifiable Information from a list of strings
    and return dict describing PII.

    Args
    texts: list of str to detect PII in
    matcher: spacy matcher object specifying the patterns to look for
    nlp: spacy language model
    output (default='counts'): the type of output to produce. Options are
        'counts': the estimated counts of PII in each of the strs
        'breakdown': the breakdown of the counts of each type of PII in each of the strs
    random_sample (default='None'): if we just want to take a sample of the string in texts to check for PII

    Return:
    dict with key, value pair for each str in texts.
    """
    PII = dict()
    n = 0

    # define the mapping between the entity types and the corresponding check digit algorithm
    # most of these do not involve check digit algorithms, so we just use the identity function
    # no_check_digit
    check_digit_algorithms = {
        "PERSON_NAME": no_check_digit,
        "PHONE_AU": no_check_digit,
        "EMAIL": no_check_digit,
        "URL": no_check_digit,
        "MEDICARE_NUMBER": candidate_medicare_valid,
        "TAX_FILE_NUMBER": candidate_tfn_valid,
        "POSTCODE": no_check_digit,
        "DATE": no_check_digit,
        "CAR_REG_VIC": no_check_digit,
        "CREDIT_CARD_NUMBER": candidate_creditcard_valid,
    }

    # check to see if there is an NLP type in pii_types
    # in which case use a language model
    # otherwise all the types are rule-based
    # this check speeds things up when language model not needed
    use_language_model = False

    nlp_types = [pii_type for pii_type in pii_types if is_nlp_type(pii_type)]
    rule_types = [pii_type for pii_type in pii_types if is_nlp_type(pii_type) == False]

    if len(nlp_types) > 0:
        use_language_model = True

    # run through all the texts and check for PII types
    for doc in nlp.pipe(texts, n_process=n_process, batch_size=batch_size):
        # dict for all pii within a given str
        item_pii = dict()
        for pii_type in pii_types:
            item_pii[pii_type] = []

        matches = matcher(doc)

        # pii types that require nlp language model
        for nlp_type in nlp_types:
            item_pii[nlp_type] = [
                ent.text
                for ent in doc.ents
                if ent.label_ == ENTITY_PII_MAPPING[nlp_type]
            ]

        # pii types defined by rules
        for match_id, start, end in matches:
            string_id = nlp.vocab.strings[match_id]
            for pii_type in rule_types:
                # check each of the pii types and
                if string_id == ENTITY_PII_MAPPING[pii_type]:
                    # apply check digit algorithm if necessary
                    check_digit_function = check_digit_algorithms[pii_type]
                    if check_digit_function(doc[start:end]):
                        # append data to appropriate entity type list
                        item_pii[pii_type].append(doc[start:end].text)
                else:
                    next
        PII[n] = item_pii
        n += 1

    return PII


def pii_from_texts_with_spans(texts, matcher, nlp, pii_types=DEFAULT_PII_TYPES):
    """
    Extract Personally Identifiable Information from a list of strings
    and return dict describing PII.
    Args
    texts: list of str to detect PII in
    matcher: spacy matcher object specifying the patterns to look for
    nlp: spacy language model
    output (default='counts'): the type of output to produce. Options are
        'counts': the estimated counts of PII in each of the strs
        'breakdown': the breakdown of the counts of each type of PII in each of the strs
    random_sample (default='None'): if we just want to take a sample of the string in texts to check for PII

    Return:
    list with key, value pair for each str in texts.
    """
    output = []
    # define the mapping between the entity types and the corresponding check digit algorithm
    # most of these do not involve check digit algorithms, so we just use the identity function
    # no_check_digit
    check_digit_algorithms = {
        "PERSON_NAME": no_check_digit,
        "PHONE_AU": no_check_digit,
        "EMAIL": no_check_digit,
        "URL": no_check_digit,
        "MEDICARE_NUMBER": candidate_medicare_valid,
        "TAX_FILE_NUMBER": candidate_tfn_valid,
        "POSTCODE": no_check_digit,
        "DATE": no_check_digit,
        "CAR_REG_VIC": no_check_digit,
        "CREDIT_CARD_NUMBER": candidate_creditcard_valid,
    }
    # check to see if there is an NLP type in pii_types
    # in which case use a language model
    # otherwise all the types are rule-based
    # this check speeds things up when language model not needed
    use_language_model = False

    nlp_types = [pii_type for pii_type in pii_types if is_nlp_type(pii_type)]
    rule_types = [pii_type for pii_type in pii_types if is_nlp_type(pii_type) == False]

    if len(nlp_types) > 0:
        use_language_model = True

    # run through all the texts and check for PII types
    for doc in nlp.pipe(texts):
        # dict for all pii within a given str
        item_pii = dict()
        for pii_type in pii_types:
            item_pii[pii_type] = []

        matches = matcher(doc)

        # pii types that require nlp language model
        for nlp_type in nlp_types:
            item_pii[nlp_type] = [
                ((ent.start_char, ent.end_char), ent.text)
                for ent in doc.ents
                if ent.label_ == ENTITY_PII_MAPPING[nlp_type]
            ]

        # pii types defined by rules
        for match_id, start, end in matches:
            string_id = nlp.vocab.strings[match_id]
            for pii_type in rule_types:
                # check each of the pii types and
                if string_id == ENTITY_PII_MAPPING[pii_type]:
                    # apply check digit algorithm if necessary
                    check_digit_function = check_digit_algorithms[pii_type]
                    if check_digit_function(doc[start:end]):
                        # append data to appropriate entity type list
                        item_pii[pii_type].append(
                            (
                                (doc[start:end].start_char, doc[start:end].end_char),
                                doc[start:end].text,
                            )
                        )  # change to character start and end
                else:
                    next
        output.append(item_pii)
    return output


def pii_summary(pii):
    """
    Produce summary of the counts of each of the PII types given a dict
    that describes the PII for each of the entries in a list of strs
    """
    summary_counts = dict()

    for row in pii.keys():
        out = pii[row]
        counts = {key: len(value) for key, value in out.items()}
        summary_counts[row] = counts

    return summary_counts


def pii_profile(df_input, matcher):
    """
    Produce a nicely formatted summary dataframe,
    showing the types of PII in each of the columns which have PII
    """
    columns = df_input.columns

    # only the text columns
    text_columns = [column for column in columns if df_input[column].dtype == "O"]

    df_pii_profile = pd.DataFrame()

    for column in text_columns:
        print(column)
        input_list = df_input[column].values
        PII = pii_from_texts(input_list, matcher)
        summary_counts = pd.DataFrame(pii_summary(PII)).T.sum()
        df_pii_profile[column] = summary_counts

    return df_pii_profile


def calculate_overall_counts(pii_results):
    """
    Given a dict describing the pii outputs from a dataframe, calculate the overall
    counts of each PII type.
    """
    columns = pii_results.keys()
    df_pii_summary = pd.DataFrame()

    for column in columns:
        PII = pii_results[column]
        summary_counts = pd.DataFrame(pii_summary(PII)).T.sum()
        df_pii_summary[column] = summary_counts

    return df_pii_summary


def calculate_overall_counts(pii_results):
    """
    Given a dict describing the pii outputs from a dataframe, calculate the overall
    values of each PII type.
    """
    columns = pii_results.keys()
    dfs = []

    for column in columns:
        PII = pii_results[column]
        column_df = pd.DataFrame(PII).apply(lambda row: row.sum(), axis=1)
        dfs.append(column_df)

    # sum up all the dataframes
    df_final = dfs[0]
    for n in range(1, len(dfs)):
        df_final += dfs[n]

    # remove the duplicate values
    df_final = df_final.apply(lambda l: list(set(l)))

    return df_final


def calculate_cell_counts(pii_results):
    """
    Given a dict describing the pii outputs from a dataframe, calculate the
    counts of each PII type within each cell.
    """
    columns = pii_results.keys()
    df_pii_summary = pd.DataFrame()

    for column in columns:
        PII = pii_results[column]
        summary_counts = (
            pd.DataFrame(pii_summary(PII)).applymap(lambda cell: len(cell)).to_dict()
        )
        df_pii_summary[column] = summary_counts

    return df_pii_summary


def inspect_for_pii(df_input, matcher, nlp):
    columns = df_input.columns

    # only the text columns
    text_columns = [column for column in columns if df_input[column].dtype == "O"]

    pii_results = dict()

    for column in text_columns:
        input_list = df_input[column].values  # only need the distinct ones
        PII = pii_from_texts(input_list, matcher, nlp)
        pii_results[column] = PII

    return pii_results


def get_pii_rows(pii_results, pii_type):
    """
    Given a results dict with the pii results in each of the rows
    return the rows indices that contain each pii type

    Return:
    list of row indices where the specified pii type appearss
    """

    row_indices = []

    for column in pii_results.keys():
        for row_index in pii_results[column].keys():
            if len(pii_results[column][row_index][pii_type]) > 0:
                row_indices.append(row_index)

    return list(set(row_indices))
